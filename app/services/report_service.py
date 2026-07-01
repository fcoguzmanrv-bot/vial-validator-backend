"""
Generador de informe técnico Word — estructura de secciones:
  1. Portada
  2. Resumen Ejecutivo
  3. Tabla General de Observaciones
  3..N. Detalle por disciplina (Alineamiento / Perfil / Drenaje / Otros)
  N+1. Firmas
"""

import io
from collections import defaultdict
from datetime import date as _date_type

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.schemas.aashto import AASHTOObservation
from app.schemas.compare import VersionChange  # kept for backward-compat signature

# ── Paleta de colores ─────────────────────────────────────────────────────────

_BLUE_DOTD  = RGBColor(0x1F, 0x4E, 0x79)
_RED_CRIT   = RGBColor(0xC0, 0x00, 0x00)
_ORANGE_MOD = RGBColor(0xE6, 0x5C, 0x00)
_GREEN_OK   = RGBColor(0x37, 0x86, 0x10)
_GRAY_TEXT  = RGBColor(0x60, 0x60, 0x60)
_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

_BG_HEADER   = "1F4E79"
_BG_CRITICO  = "FFE7E7"
_BG_MODERADO = "FFF3E0"
_BG_CUMPLE   = "F0FFF0"
_BG_INFO     = "F8F8F8"

# Orden canónico de disciplinas
_DISCIPLINES = [
    "Alineamiento Horizontal y Sección Transversal",
    "Alzado / Perfil Vertical",
    "Drenaje",
    "Otros",
]


# ── Clasificación por disciplina ──────────────────────────────────────────────

def _classify_observation(parameter: str) -> str:
    p = parameter.lower()
    if any(k in p for k in [
        "radio", "curvatura", "compuesta", "broken", "broken-back",
        # con y sin tilde (datos de PDF pueden carecer de acentos)
        "deflexion", "deflexión",
        "peralte",
        "superelevacion", "superelevación",
        "transicion", "transición",
        "transversal",
        "angulo", "ángulo",
        "horizontal", "curva compuesta",
    ]):
        return "Alineamiento Horizontal y Sección Transversal"
    elif any(k in p for k in [
        "vertical", "pendiente longitudinal",
        "drenaje superficial", "k value",
        "curva vertical", "escurrimiento", "rasante",
    ]):
        return "Alzado / Perfil Vertical"
    elif any(k in p for k in [
        "tuberia", "tubería", "pipe", "outlet",
        "differential", "tapada",
        "hidraulica", "hidráulica",
        "drenaje —", "capacidad",
    ]):
        return "Drenaje"
    else:
        return "Otros"


# ── Referencias normativas detalladas ────────────────────────────────────────

_NORMATIVE_REFERENCES: dict[str, str] = {
    # Alineamiento Horizontal
    "radio":            "DOTD RDM Chapter 4, Sec. 4.2 (April 2022), p. 4-3 / AASHTO Green Book 7th ed., Table 3-7",
    "curvatura":        "DOTD RDM Chapter 4, Sec. 4.2 (April 2022), p. 4-3 / AASHTO Green Book 7th ed., Table 3-7",
    "compuesta":        "DOTD RDM Chapter 4, Sec. 4.2.1 (April 2022), p. 4-3",
    "broken":           "DOTD RDM Chapter 4, Sec. 4.2.1 (April 2022), p. 4-4",
    "deflexion":        "DOTD RDM Chapter 4, Sec. 4.2.1 (April 2022), p. 4-4",
    "deflexión":        "DOTD RDM Chapter 4, Sec. 4.2.1 (April 2022), p. 4-4",
    "angulo":           "DOTD RDM Chapter 4, Sec. 4.2.1 (April 2022), p. 4-4",
    "ángulo":           "DOTD RDM Chapter 4, Sec. 4.2.1 (April 2022), p. 4-4",
    "peralte invertido":"DOTD RDM Chapter 4, Sec. 4.6.1 (April 2022), p. 4-17 / AASHTO Green Book §3.3",
    "transicion":       "DOTD RDM Chapter 4, Sec. 4.6.3 (April 2022), pp. 4-18 a 4-26",
    "transición":       "DOTD RDM Chapter 4, Sec. 4.6.3 (April 2022), pp. 4-18 a 4-26",
    "superelevacion":   "DOTD RDM Chapter 4, Sec. 4.6 (April 2022), p. 4-17 / AASHTO Green Book §3.3",
    "superelevación":   "DOTD RDM Chapter 4, Sec. 4.6 (April 2022), p. 4-17 / AASHTO Green Book §3.3",
    "transversal":      "DOTD RDM Chapter 5, Sec. 5.2 (April 2022) / AASHTO Green Book §3.3",
    # Alzado / Perfil Vertical
    "vertical":              "DOTD RDM Chapter 4, Sec. 4.3 (April 2022), p. 4-6 / AASHTO Green Book Tables 3-34/3-36",
    "rasante":               "DOTD RDM Chapter 4, Sec. 4.3 (April 2022), p. 4-6",
    "escurrimiento":         "DOTD RDM Chapter 4, Sec. 4.5.1 (April 2022), p. 4-10",
    "pendiente longitudinal":"DOTD RDM Chapter 4, Sec. 4.3.1 (April 2022), p. 4-6 / DOTD Hydraulics Manual Sec. 8.5.2 (2011)",
    # Drenaje
    "tuberia":    "DOTD Hydraulics Manual (2011), Sec. 8.10.6, Table 6-A.4-1",
    "tubería":    "DOTD Hydraulics Manual (2011), Sec. 8.10.6, Table 6-A.4-1",
    "pipe":       "DOTD Hydraulics Manual (2011), Sec. 8.10.6, Table 6-A.4-1",
    "outlet":     "DOTD Hydraulics Manual (2011), Sec. 6.10/6.11, 8.10.7",
    "differential":"DOTD Hydraulics Manual (2011), Sec. 6.9.1/6.9.2",
    "hidraulica": "DOTD Hydraulics Manual (2011), Sec. 6.9, 6.10, 8.10.7",
    "hidráulica": "DOTD Hydraulics Manual (2011), Sec. 6.9, 6.10, 8.10.7",
    "tapada":     "DOTD Hydraulics Manual (2011), Sec. 6.7/8.12",
    "velocidad":  "DOTD Hydraulics Manual (2011), Sec. 8.10.6",
    "capacidad":  "DOTD Hydraulics Manual (2011), Sec. 6.9, 6.10, 8.10.7",
}


def _get_normative_reference(parameter: str) -> str:
    """Retorna la referencia normativa detallada según el tipo de observación."""
    p = parameter.lower()
    for keyword, reference in _NORMATIVE_REFERENCES.items():
        if keyword in p:
            return reference
    return "DOTD Louisiana Road Design Manual (April 2022) / AASHTO Green Book (7th ed.)"


# ── Helpers de XML ────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _remove_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _set_col_widths(table, widths_cm: list[float]) -> None:
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)


def _dxa(n: int) -> int:
    """1 DXA (twip) = 914400/1440 = 635 EMU."""
    return n * 635


def _set_col_widths_dxa(table, widths_dxa: list[int]) -> None:
    for row in table.rows:
        for i, w in enumerate(widths_dxa):
            if i < len(row.cells):
                row.cells[i].width = _dxa(w)


def _add_deco_line(doc, color_hex: str = "1F4E79") -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "8")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)


def _add_item_separator(doc) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "2")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "C0C0C0")
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)


def _add_footer(doc: Document, project_name: str) -> None:
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    pPr = p._p.get_or_add_pPr()
    tabs_el = OxmlElement("w:tabs")
    tab_right = OxmlElement("w:tab")
    tab_right.set(qn("w:val"), "right")
    tab_right.set(qn("w:pos"), "9360")
    tabs_el.append(tab_right)
    pPr.append(tabs_el)

    r_proj = p.add_run(project_name or "Informe Vial Validator")
    r_proj.font.size = Pt(8)
    r_proj.font.color.rgb = _GRAY_TEXT

    r_tab = p.add_run("\t")
    r_tab.font.size = Pt(8)

    r_pag = p.add_run("Página ")
    r_pag.font.size = Pt(8)
    r_pag.font.color.rgb = _GRAY_TEXT

    def _field(instr: str) -> None:
        for fld_type, fld_text in [("begin", None), (None, instr), ("end", None)]:
            run = p.add_run()
            run.font.size = Pt(8)
            run.font.color.rgb = _GRAY_TEXT
            if fld_type:
                el = OxmlElement("w:fldChar")
                el.set(qn("w:fldCharType"), fld_type)
            else:
                el = OxmlElement("w:instrText")
                el.set(qn("xml:space"), "preserve")
                el.text = f" {fld_text} "
            run._r.append(el)

    _field("PAGE")
    r_de = p.add_run(" de ")
    r_de.font.size = Pt(8)
    r_de.font.color.rgb = _GRAY_TEXT
    _field("NUMPAGES")


def _header_row(table, *labels: str) -> None:
    row = table.rows[0]
    for i, label in enumerate(labels):
        if i >= len(row.cells):
            break
        cell = row.cells[i]
        cell.text = label
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(cell, _BG_HEADER)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = _WHITE
        run.font.size = Pt(9)


def _row_bg(obs: AASHTOObservation) -> str:
    if obs.severity == "critico" and not obs.complies:
        return _BG_CRITICO
    if obs.severity == "moderado" and not obs.complies:
        return _BG_MODERADO
    if obs.complies:
        return _BG_CUMPLE
    return _BG_INFO


def _deduplicate(observations: list[AASHTOObservation]) -> list[AASHTOObservation]:
    seen: set[tuple[str, str]] = set()
    out: list[AASHTOObservation] = []
    for obs in observations:
        key = (obs.parameter.strip(), obs.found_value.strip())
        if key not in seen:
            seen.add(key)
            out.append(obs)
    return out


def _first_sentence(text: str, max_len: int = 80) -> str:
    """Primera oración sin cortar en puntos decimales.
    Busca '. ' (punto + espacio) como separador de oración."""
    idx = text.find(". ")
    if 0 < idx <= max_len:
        return text[:idx]
    return text[:max_len]


# ── Secciones ─────────────────────────────────────────────────────────────────

def _section_portada(
    doc: Document,
    project_name: str,
    contract_number: str,
    engineer: str,
    reviewing_firm: str,
    page_range: str,
    pdf_filename: str,
    today: str,
) -> None:
    for _ in range(5):
        doc.add_paragraph()

    _add_deco_line(doc)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(8)
    p_title.paragraph_format.space_after  = Pt(4)
    r = p_title.add_run("INFORME TÉCNICO DE VALIDACIÓN GEOMÉTRICA")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = _BLUE_DOTD

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after  = Pt(8)
    r2 = p_sub.add_run("Cumplimiento Normativo DOTD Louisiana Road Design Manual")
    r2.font.size = Pt(12)
    r2.font.color.rgb = _GRAY_TEXT

    _add_deco_line(doc)
    doc.add_paragraph()

    tbl = doc.add_table(rows=4, cols=2)
    _remove_table_borders(tbl)
    _set_col_widths(tbl, [8.5, 8.5])

    def _cell(row_idx: int, col_idx: int, label: str, value: str) -> None:
        cell = tbl.rows[row_idx].cells[col_idx]
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        rl = p.add_run(f"{label}: ")
        rl.bold = True
        rl.font.size = Pt(10)
        rl.font.color.rgb = _BLUE_DOTD
        rv = p.add_run(value or "—")
        rv.font.size = Pt(10)

    _cell(0, 0, "Proyecto",              project_name)
    _cell(0, 1, "Contrato N°",           contract_number)
    _cell(1, 0, "Planos / Páginas",
          f"{pdf_filename}  |  págs. {page_range}" if pdf_filename else page_range)
    _cell(1, 1, "Fecha",                 today)
    _cell(2, 0, "Ingeniero responsable", engineer)
    _cell(2, 1, "Empresa revisora",      reviewing_firm)

    merged = tbl.rows[3].cells[0].merge(tbl.rows[3].cells[1])
    merged.paragraphs[0].clear()
    p4 = merged.paragraphs[0]
    p4.paragraph_format.space_before = Pt(3)
    p4.paragraph_format.space_after  = Pt(3)
    rl4 = p4.add_run("Normativa aplicada: ")
    rl4.bold = True
    rl4.font.size = Pt(10)
    rl4.font.color.rgb = _BLUE_DOTD
    rv4 = p4.add_run(
        "DOTD Louisiana Road Design Manual (April 2022) / AASHTO Green Book (7th ed.)"
    )
    rv4.font.size = Pt(10)

    doc.add_paragraph()
    _add_deco_line(doc)


def _section_resumen(
    doc: Document,
    project_name: str,
    pdf_filename: str,
    page_range: str,
    all_obs: list[AASHTOObservation],
) -> None:
    non_compliant = [o for o in all_obs if not o.complies]
    criticos      = [o for o in non_compliant if o.severity == "critico"]
    moderados     = [o for o in non_compliant if o.severity == "moderado"]
    informativos  = [o for o in all_obs if o.complies and o.severity == "informativo"]
    n_total       = len(non_compliant)
    n_critico     = len(criticos)
    n_moderado    = len(moderados)
    n_info        = len(informativos)

    h1 = doc.add_heading("1. Resumen Ejecutivo", level=1)
    h1.runs[0].font.color.rgb = _BLUE_DOTD

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.add_run(
        f"El presente informe documenta los resultados de la validación geométrica del proyecto "
        f"{project_name or '[Proyecto]'}, realizada sobre las páginas {page_range or '[rango]'} "
        f"del juego de planos {pdf_filename or '[plano]'}. Se identificaron {n_total} incumplimientos "
        f"normativos bajo el estándar DOTD Louisiana Road Design Manual (April 2022) y AASHTO Green "
        f"Book (7th ed.), de los cuales {n_critico} son de severidad crítica y {n_moderado} moderada."
    ).font.size = Pt(10)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    _header_row(tbl, "Métrica", "Valor")
    _set_col_widths(tbl, [10.5, 5.0])

    metrics = [
        ("Páginas analizadas",                                  page_range or "—"),
        ("Total incumplimientos",                               str(n_total)),
        ("Críticos",                                            str(n_critico)),
        ("Moderados",                                           str(n_moderado)),
        ("Observaciones informativas (excluidas del informe)",  str(n_info)),
    ]
    for label, val in metrics:
        row = tbl.add_row()
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9.5)
        row.cells[1].text = val
        row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9.5)

    doc.add_paragraph()

    if criticos:
        h2 = doc.add_heading("Hallazgos Críticos Destacados", level=2)
        h2.runs[0].font.color.rgb = _RED_CRIT

        for obs in criticos[:5]:
            p_item = doc.add_paragraph()
            p_item.paragraph_format.space_before = Pt(1)
            p_item.paragraph_format.space_after  = Pt(1)
            param_short  = obs.parameter[:50]
            frag = _first_sentence(obs.observation or obs.found_value or "", 80)
            r = p_item.add_run(f"• {param_short}: {frag}")
            r.font.size = Pt(9.5)


def _section_tabla_general(
    doc: Document,
    all_obs: list[AASHTOObservation],
) -> None:
    h1 = doc.add_heading("2. Observaciones DOTD / AASHTO", level=1)
    h1.runs[0].font.color.rgb = _BLUE_DOTD

    cols     = ["Parámetro", "Valor encontrado", "Valor normativo", "Cumple", "Observación"]
    col_dxa  = [2200, 1400, 1800, 400, 3560]   # total = 9360 DXA

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    _header_row(tbl, *cols)
    _set_col_widths_dxa(tbl, col_dxa)

    for obs in all_obs:
        row   = tbl.add_row()
        bg    = _row_bg(obs)
        check = "✓" if obs.complies else "✗"
        obs_text = (obs.observation or "")
        if len(obs_text) > 150:
            obs_text = obs_text[:147] + "..."

        values = [obs.parameter, obs.found_value, obs.normative_value, check, obs_text]
        check_colors = {3: _GREEN_OK if obs.complies else _RED_CRIT}

        for i, val in enumerate(values):
            cell = row.cells[i]
            cell.width = _dxa(col_dxa[i])   # garantiza 9360 DXA en cada fila de datos
            _set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            if i == 3:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(val)
            run.font.size = Pt(9)           # fuente 9pt en toda la tabla
            if i in check_colors:
                run.font.color.rgb = check_colors[i]
                run.bold = True

    doc.add_paragraph()


def _render_obs_items(
    doc: Document,
    obs_list: list[AASHTOObservation],
    color: RGBColor,
) -> None:
    """Renderiza la lista detallada de observaciones sin encabezado de sección."""
    for i, obs in enumerate(obs_list, 1):
        h3 = doc.add_heading(level=3)
        h3.clear()
        r_num = h3.add_run(f"{i}. {obs.parameter}")
        r_num.font.color.rgb = color
        r_num.font.size = Pt(10)

        def _detail_line(bold_label: str, text: str, color_run: RGBColor | None = None) -> None:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            rl = p.add_run(f"{bold_label}: ")
            rl.bold = True
            rl.font.size = Pt(9.5)
            rv = p.add_run(text or "—")
            rv.font.size = Pt(9.5)
            if color_run:
                rv.font.color.rgb = color_run

        _detail_line("Valor encontrado", obs.found_value)
        _detail_line("Valor normativo",  obs.normative_value)
        _detail_line("Observación",      obs.observation or "—", color_run=color)

        p_ref = doc.add_paragraph()
        p_ref.paragraph_format.space_before = Pt(2)
        p_ref.paragraph_format.space_after  = Pt(4)
        r_ref = p_ref.add_run(f"Referencia: {_get_normative_reference(obs.parameter)}")
        r_ref.italic = True
        r_ref.font.size = Pt(8.5)
        r_ref.font.color.rgb = _GRAY_TEXT

        _add_item_separator(doc)


def _section_detalle_por_disciplina(
    doc: Document,
    non_compliant: list[AASHTOObservation],
    start_num: int,
) -> int:
    """Genera secciones de detalle agrupadas por disciplina.
    Retorna el número de sección siguiente (para Firmas)."""
    by_disc: dict[str, list[AASHTOObservation]] = defaultdict(list)
    for obs in non_compliant:
        by_disc[_classify_observation(obs.parameter)].append(obs)

    sec = start_num
    for discipline in _DISCIPLINES:
        disc_obs = by_disc.get(discipline, [])
        if not disc_obs:
            continue

        criticos  = [o for o in disc_obs if o.severity == "critico"]
        moderados = [o for o in disc_obs if o.severity == "moderado"]

        h1 = doc.add_heading(f"{sec}. {discipline}", level=1)
        h1.runs[0].font.color.rgb = _BLUE_DOTD

        if criticos:
            h2 = doc.add_heading(f"{sec}.1 Incumplimientos Críticos", level=2)
            h2.runs[0].font.color.rgb = _RED_CRIT
            _render_obs_items(doc, criticos, _RED_CRIT)

        if moderados:
            sub = 2 if criticos else 1
            h2 = doc.add_heading(f"{sec}.{sub} Incumplimientos Moderados", level=2)
            h2.runs[0].font.color.rgb = _ORANGE_MOD
            _render_obs_items(doc, moderados, _ORANGE_MOD)

        sec += 1

    return sec


def _section_firmas(doc: Document, engineer: str, today: str, section_num: int = 6) -> None:
    h1 = doc.add_heading(f"{section_num}. Firmas", level=1)
    h1.runs[0].font.color.rgb = _BLUE_DOTD

    tbl = doc.add_table(rows=2, cols=2)
    _remove_table_borders(tbl)
    _set_col_widths(tbl, [8.5, 8.5])

    for col_idx, (header, nombre_val) in enumerate([
        ("Preparado por", f"Nombre: {engineer or ''}"),
        ("Revisado por",  "Nombre:"),
    ]):
        cell_hdr = tbl.rows[0].cells[col_idx]
        cell_hdr.paragraphs[0].clear()
        p_hdr = cell_hdr.paragraphs[0]
        p_hdr.paragraph_format.space_before = Pt(4)
        p_hdr.paragraph_format.space_after  = Pt(2)
        r = p_hdr.add_run(header)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = _BLUE_DOTD

        cell_body = tbl.rows[1].cells[col_idx]
        cell_body.paragraphs[0].clear()
        p_firma = cell_body.paragraphs[0]
        p_firma.paragraph_format.space_before = Pt(20)
        p_firma.paragraph_format.space_after  = Pt(2)
        r_f = p_firma.add_run("Firma: " + "_" * 50)
        r_f.font.size = Pt(10)

        p_nombre = cell_body.add_paragraph()
        p_nombre.paragraph_format.space_before = Pt(6)
        p_nombre.paragraph_format.space_after  = Pt(2)
        p_nombre.add_run(nombre_val).font.size = Pt(10)

        fecha_val = f"Fecha: {today}" if col_idx == 0 else "Fecha:"
        p_fecha = cell_body.add_paragraph()
        p_fecha.paragraph_format.space_before = Pt(3)
        p_fecha.paragraph_format.space_after  = Pt(2)
        p_fecha.add_run(fecha_val).font.size = Pt(10)


# ── Punto de entrada público ──────────────────────────────────────────────────

def build_report(
    observations: list[AASHTOObservation],
    project_name: str = "",
    contract_number: str = "",
    engineer: str = "",
    reviewing_firm: str = "",
    page_range: str = "",
    pdf_filename: str = "",
    changes: "list[VersionChange] | None" = None,
    responsible_engineer: str = "",   # alias de engineer (compat con endpoint)
    report_date: str = "",
) -> bytes:
    # Compatibilidad: el endpoint pasa engineer= directamente; responsible_engineer es fallback
    if responsible_engineer and not engineer:
        engineer = responsible_engineer
    today = report_date or _date_type.today().isoformat()

    doc = Document()

    section = doc.sections[0]
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    _add_footer(doc, project_name)

    deduped = _deduplicate(observations)
    non_compliant = [o for o in deduped if not o.complies]

    # ── Portada ────────────────────────────────────────────────────────────────
    _section_portada(doc, project_name, contract_number, engineer, reviewing_firm,
                     page_range, pdf_filename, today)
    doc.add_page_break()

    # ── 1. Resumen Ejecutivo ──────────────────────────────────────────────────
    _section_resumen(doc, project_name, pdf_filename, page_range, deduped)

    # ── 2. Tabla General ──────────────────────────────────────────────────────
    _section_tabla_general(doc, deduped)

    # ── 3..N: Detalle por disciplina ──────────────────────────────────────────
    next_sec = _section_detalle_por_disciplina(doc, non_compliant, start_num=3)

    # ── Firmas ────────────────────────────────────────────────────────────────
    doc.add_page_break()
    _section_firmas(doc, engineer, today, section_num=next_sec)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
